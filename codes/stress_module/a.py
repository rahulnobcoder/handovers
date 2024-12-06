from docx import Document
from docx.shared import Pt

# Create a new document
doc = Document()

# Define the data from the image into a table format
data = [
    ["Direction", "DISFA → BP4D", "", "", "", "Avg.", "BP4D → DISFA", "", "", "", "Avg."],
    ["AU", "1", "2", "4", "6", "", "1", "2", "4", "6", ""],
    ["DRML [36]", "19.4", "16.9", "22.4", "58.0", "64.5", "36.3", "10.4", "7.0", "16.9", "14.1"],
    ["JÂA-Net [25]", "10.9", "6.7", "42.4", "52.9", "68.3", "36.2", "12.5", "13.2", "27.6", "23.8"],
    ["ME-GraphAU [21]", "36.5", "30.3", "35.8", "48.8", "62.2", "42.7", "43.3", "22.5", "38.4", "33.1"],
    ["ME-GraphAU + FFHQ pre-train", "20.1", "32.9", "38.0", "64.0", "73.0", "30.6", "34.4", "14.4", "54.4", "33.7"],
    ["GH-Feat [31]", "29.4", "30.0", "37.1", "64.0", "73.5", "18.9", "15.2", "27.5", "52.7", "32.9"],
    ["Patch-MCD* [32]", "-", "-", "-", "-", "-", "34.3", "16.6", "52.1", "33.5", "37.4"],
    ["IdenNet* [29]", "-", "-", "-", "-", "-", "20.1", "25.5", "37.3", "49.6", "39.7"],
    ["Ours", "51.4", "46.0", "36.0", "49.6", "49.0", "61.3", "70.5", "36.3", "42.2", "61.5", "54.4"]
]

# Ensure that every row has the correct number of columns (12 columns)
for row in data:
    while len(row) < 12:
        row.append("")  # Add empty strings to match the column count

# Add table to the document
table = doc.add_table(rows=1, cols=len(data[0]))

# Set table headers
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(data[0]):
    hdr_cells[i].text = column_name

# Add data rows to the table
for row_data in data[1:]:
    row = table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row[i].text = cell_data

# Style the table (optional)
table.style = 'Table Grid'
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# Save the document
file_path = "converted_table_fixed.docx"
doc.save(file_path)

print(f"Document saved as {file_path}")
